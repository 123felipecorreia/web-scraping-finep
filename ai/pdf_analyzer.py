import os
import sys
import re
from typing import Dict, List
import io
import urllib.parse as urlparse

# Optional third-party helpers used for text extraction
try:
    from bs4 import BeautifulSoup
except Exception:
    BeautifulSoup = None

try:
    import trafilatura
except Exception:
    trafilatura = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import docx
except Exception:
    docx = None
# Ensure project root is on sys.path when running this file directly (python ai/pdf_analyzer.py)
_CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(_CURRENT_DIR)
if _PROJECT_ROOT not in sys.path:
    sys.path.insert(0, _PROJECT_ROOT)

from ai.openai_client import OpenAIClient
from utils.text_processor import split_text_intelligently, safe_json_parse

class PDFAnalyzer:
    def __init__(self):
        self.openai_client = OpenAIClient()

    # ---- small helper utilities (moved here per request) ----
    def _guess_ext(self, url: str) -> str:
        path = urlparse.urlparse(url).path
        ext = os.path.splitext(path)[1].lower()
        return ext

    def _extract_text_from_pdf(self, content: bytes) -> str:
        if fitz is None:
            return ""
        try:
            with fitz.open(stream=content, filetype="pdf") as doc:
                texts = []
                for page in doc:
                    texts.append(page.get_text("text"))
                return "\n".join(texts)
        except Exception:
            return ""

    def _extract_text_from_docx(self, content: bytes) -> str:
        if docx is None:
            return ""
        try:
            mem = io.BytesIO(content)
            d = docx.Document(mem)
            return "\n".join([p.text for p in d.paragraphs])
        except Exception:
            return ""

    def _extract_text_from_html(self, html: str, url: str) -> str:
        if trafilatura is not None:
            try:
                downloaded = trafilatura.extract(html, include_links=False, include_formatting=False, url=url)
                if downloaded and downloaded.strip():
                    return downloaded
            except Exception:
                pass
        if BeautifulSoup is None:
            return html
        soup = BeautifulSoup(html, "lxml")
        return soup.get_text(separator="\n")
    def analyze_pdf_with_ai_robust(self, pdf_text: str) -> Dict:
        """Análise de PDF otimizada"""
        if not self.openai_client.is_available():
            return {
                'Valor_Global': "IA não disponível",
                'Valor_Maximo_Por_Projeto': "IA não disponível", 
                'Data_Limite_Submissao': "IA não disponível",
                'Percentual_Contrapartida': "IA não disponível",
                'Nivel_TRL_Exigido': "IA não disponível"
            }
        
        try:
            text_chunks = split_text_intelligently(pdf_text, 6000)
            print(f"   📖 Analisando {len(text_chunks)} seções do PDF...")
            
            all_results = []
            informacoes_encontradas = []
            
            for i, chunk in enumerate(text_chunks, 1):
                print(f"      📄 Seção {i}/{len(text_chunks)}...", end=" ")
                
                prompt = self._create_analysis_prompt(chunk)
                result = self._process_chunk_with_ai(prompt)
                
                if result:
                    all_results.append(result)
                    found_info = self._extract_found_info(result)
                    if found_info:
                        print("✅")
                        for info in found_info:
                            informacoes_encontradas.append(info)
                            print(f"         🎯 {info}")
                    else:
                        print("❌")
                else:
                    print("❌")
            
            self._display_extraction_summary(informacoes_encontradas)
            return self._consolidate_results(all_results)
            
        except Exception as e:
            print(f"\n   ❌ Erro geral na análise: {e}")
            return self._get_error_result(str(e))

    from typing import Tuple

    def is_pdf_edital(self, pdf_text: str, pdf_url: str = "", project_name: str = "") -> Tuple[bool, float]:
        """Use LLM (when available) to judge whether a PDF text is the official edital/chamada.
        Returns (is_edital: bool, confidence: float 0..1).
        MELHORADO: Agora prioriza editais RECENTES (2024-2025).
        """
        text = (pdf_text or "").strip()
        
        # Keywords atualizadas para focar em editais recentes
        keywords_positive = [
            'edital', 'chamada pública', 'chamada', 'objeto', 'submissão', 'proposta', 'cronograma',
            'recursos', 'verba', 'financiamento', 'seleção', 'avaliação', 'proponente', 'inscrição',
            'nº do edital', 'número do edital', 'regulamento', 'concurso', '2024', '2025'
        ]
        keywords_negative = [
            'tutorial', 'como acessar', 'guia', 'manual', 'passo a passo', 'login', 'acesso', 'tutorial de',
            'como usar', 'instalação', 'suporte técnico', 'erro', 'faq', 'resultado final', 'homologado',
            'encerrado', '2023', '2022', '2021', 'arquivado'
        ]

        low_text = text.lower()
        pos_count = sum(1 for k in keywords_positive if k in low_text)
        neg_count = sum(1 for k in keywords_negative if k in low_text)

        filename = ""
        try:
            p = urlparse.urlparse(pdf_url)
            filename = os.path.basename(p.path or "")
        except Exception:
            filename = ""

        if project_name:
            pn = project_name.lower()
            if pn in filename.lower() or pn in low_text:
                pos_count += 2

        # MELHORADO: Prioriza anos 2024 e 2025
        try:
            from datetime import datetime
            current_year = str(datetime.utcnow().year)
            next_year = str(datetime.utcnow().year + 1)
            
            # Boost maior para ano atual e próximo
            if current_year in filename or current_year in pdf_url or current_year in low_text:
                pos_count += 3
            if next_year in filename or next_year in pdf_url or next_year in low_text:
                pos_count += 2
                
            # Penaliza anos antigos
            old_years = ['2023', '2022', '2021', '2020']
            for old_year in old_years:
                if old_year in filename or old_year in pdf_url or old_year in low_text:
                    neg_count += 2
                
        except Exception:
            pass

        heuristic_conf = max(0.0, (pos_count - neg_count) / max(len(keywords_positive), 1))
        
        if self.openai_client.is_available() and len(text) > 200:
            try:
                # PROMPT MELHORADO: Foco em editais recentes e ativos
                prompt = f"""
Classificador de Editais de Fomento (Português) - PRIORIDADE: EDITAIS RECENTES E ATIVOS

Você deve analisar se este documento é um EDITAL/CHAMADA PÚBLICA OFICIAL e RECENTE (2024-2025) para submissão de propostas.

CRITÉRIOS DE ALTA PRIORIDADE (aumentam confiança):
1. Documento menciona explicitamente anos 2024 ou 2025
2. Contém prazos de submissão FUTUROS ou VIGENTES
3. Usa termos: "em aberto", "prazo em vigor", "inscrições abertas"
4. Apresenta estrutura completa: Objeto + Valor + Prazo + Elegibilidade + Critérios

CRITÉRIOS DE EXCLUSÃO (reduzem confiança drasticamente):
1. Documento menciona anos 2023, 2022, 2021 ou anteriores
2. Contém termos: "encerrado", "finalizado", "resultado homologado", "arquivado"
3. É tutorial, guia, manual de sistema ou FAQ
4. Prazos já vencidos ou no passado

ANÁLISE DE VIGÊNCIA:
- Se houver datas, verifique se são FUTURAS (após {datetime.now().strftime('%m/%Y')})
- Editais com prazos vencidos devem ter confidence < 0.3

URL do documento: {pdf_url}

TRECHO DO TEXTO (primeiros 8000 caracteres):
{text[:8000]}

Responda APENAS com JSON válido:
{{"is_edital": true|false, "confidence": 0.0-1.0, "year_detected": "YYYY ou null", "status": "ativo|encerrado|indefinido"}}

EXEMPLOS:

Exemplo 1 - EDITAL ATIVO 2025 (alta confiança):
{{"is_edital": true, "confidence": 0.95, "year_detected": "2025", "status": "ativo"}}

Exemplo 2 - EDITAL ENCERRADO 2023 (baixa confiança):
{{"is_edital": true, "confidence": 0.25, "year_detected": "2023", "status": "encerrado"}}

Exemplo 3 - TUTORIAL (rejeitar):
{{"is_edital": false, "confidence": 0.05, "year_detected": null, "status": "indefinido"}}
"""

                resp = self.openai_client.client.chat.completions.create(
                    model="gpt-4o-mini" if hasattr(self.openai_client.client, 'chat') else "gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "Você é um classificador especializado em identificar editais de fomento RECENTES e ATIVOS. Priorize documentos de 2024-2025."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.0,
                    max_tokens=150,
                )
                raw = resp.choices[0].message.content.strip()
                parsed = safe_json_parse(raw)
                
                if isinstance(parsed, dict) and 'is_edital' in parsed:
                    is_ed = bool(parsed.get('is_edital'))
                    conf = float(parsed.get('confidence') or 0.0)
                    status = parsed.get('status', 'indefinido')
                    
                    # Penaliza editais encerrados
                    if status == 'encerrado':
                        conf = conf * 0.3
                    
                    combined = (heuristic_conf + conf) / 2.0
                    return is_ed, max(0.0, min(1.0, combined))
            except Exception:
                pass

        # Fallback: treat heuristic_conf > 0.2 as positive
        return (heuristic_conf > 0.2), float(heuristic_conf)
    
    def _create_analysis_prompt(self, chunk: str) -> str:
        """PROMPT MELHORADO: Extração precisa com foco em editais recentes"""
        current_year = datetime.now().year
        return f"""
SISTEMA DE EXTRAÇÃO DE DADOS DE EDITAIS DE FOMENTO - ANO FOCAL: {current_year}/{current_year+1}

MISSÃO: Extrair informações PRECISAS e VERIFICÁVEIS de chamadas públicas/editais de fomento à pesquisa e inovação.

CAMPOS OBRIGATÓRIOS A EXTRAIR:

1. **Valor Global** (Recursos Totais)
   - Procure: "valor global", "recursos disponíveis", "orçamento total", "dotação orçamentária"
   - Formato esperado: "R$ X.XXX.XXX,XX" ou "X milhões de reais"
   - Se não encontrar: retorne "Não encontrado"

2. **Valor Máximo por Projeto**
   - Procure: "valor máximo por projeto", "teto de financiamento", "limite por proposta"
   - Formato esperado: "R$ XXX.XXX,XX" ou "até R$ XXX mil"
   - Se não encontrar: retorne "Não encontrado"

3. **Data Limite de Submissão** ⚠️ CRÍTICO: Deve ser data FUTURA
   - Procure: "data limite", "prazo final", "até", "encerramento das inscrições"
   - Formato esperado: "DD/MM/YYYY" ou "DD de Mês de YYYY"
   - VALIDAÇÃO: Se a data for de {current_year-1} ou anterior, retorne "Prazo vencido"
   - Se não encontrar: retorne "Não encontrado"

4. **Percentual de Contrapartida**
   - Procure: "contrapartida", "aporte institucional", "recursos próprios exigidos"
   - Formato esperado: "X%" ou "X por cento"
   - Se não houver contrapartida: retorne "Não exigida"
   - Se não encontrar: retorne "Não encontrado"

5. **Nível TRL Exigido** (Technology Readiness Level)
   - Procure: "TRL", "maturidade tecnológica", "nível de prontidão"
   - Formato esperado: "TRL X", "TRL X-Y", ou inferência baseada em:
     * Pesquisa básica → "TRL 1-2 (inferido)"
     * Prova de conceito → "TRL 3-4 (inferido)"
     * Protótipo → "TRL 5-6 (inferido)"
     * Piloto/Demonstração → "TRL 7-8 (inferido)"
     * Comercialização → "TRL 9 (inferido)"
   - Se não encontrar e não puder inferir: retorne "Não encontrado"

REGRAS DE EXTRAÇÃO:
- Use APENAS informações do texto fornecido
- NÃO invente ou assuma valores
- Se houver múltiplos valores, escolha o mais recente ou o mais específico
- Mantenha números e formatação originais quando possível
- Para datas, normalize para formato DD/MM/YYYY

TEXTO PARA ANÁLISE:
```
{chunk}
```

RESPONDA ESTRITAMENTE NESTE FORMATO JSON (sem texto adicional):
{{"Valor_Global": "valor ou Não encontrado", "Valor_Maximo_Por_Projeto": "valor ou Não encontrado", "Data_Limite_Submissao": "data ou Não encontrado ou Prazo vencido", "Percentual_Contrapartida": "percentual ou Não exigida ou Não encontrado", "Nivel_TRL_Exigido": "TRL ou Não encontrado"}}
"""
    
    def _process_chunk_with_ai(self, prompt: str) -> Dict:
        """Processa um chunk com IA"""
        try:
            response = self.openai_client.client.chat.completions.create(
                model="gpt-4.0-turbo",
                messages=[
                    {"role": "system", "content": "Retorne APENAS JSON válido."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=150,
                temperature=0.05
            )
            
            ai_response = response.choices[0].message.content.strip()
            return self._parse_ai_response(ai_response)
            
        except Exception:
            return {}
    
    def _parse_ai_response(self, ai_response: str) -> Dict:
        """Parse da resposta da IA"""
        result = {}
        patterns = {
            'Valor_Global': r'"?Valor_Global"?\s*:\s*"([^"]*)"',
            'Valor_Maximo_Por_Projeto': r'"?Valor_Maximo_Por_Projeto"?\s*:\s*"([^"]*)"',
            'Data_Limite_Submissao': r'"?Data_Limite_Submissao"?\s*:\s*"([^"]*)"',
            'Percentual_Contrapartida': r'"?Percentual_Contrapartida"?\s*:\s*"([^"]*)"',
            'Nivel_TRL_Exigido': r'"?Nivel_TRL_Exigido"?\s*:\s*"([^"]*)"'
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, ai_response, re.IGNORECASE)
            if match:
                result[key] = match.group(1)
            else:
                result[key] = "Não encontrado"
        
        return result
    
    def _extract_found_info(self, result: Dict) -> List[str]:
        """Extrai informações encontradas para display"""
        found_info = []
        for key, value in result.items():
            if value != "Não encontrado":
                found_info.append(f"{key}: {value}")
        return found_info
    
    def _display_extraction_summary(self, informacoes_encontradas: List[str]):
        """Exibe resumo da extração"""
        print(f"\n   📊 INFORMAÇÕES EXTRAÍDAS DO PDF:")
        if informacoes_encontradas:
            for info in set(informacoes_encontradas):
                print(f"      ✅ {info}")
        else:
            print("      ⚠️ Nenhuma informação específica extraída")
    
    def _consolidate_results(self, all_results: List[Dict]) -> Dict:
        """Consolida resultados de todos os chunks"""
        if all_results:
            consolidated = {
                'Valor_Global': "Não encontrado",
                'Valor_Maximo_Por_Projeto': "Não encontrado",
                'Data_Limite_Submissao': "Não encontrado",
                'Percentual_Contrapartida': "Não encontrado",
                'Nivel_TRL_Exigido': "Não encontrado"
            }
            
            for key in consolidated.keys():
                for result in all_results:
                    if key in result and result[key] != "Não encontrado":
                        consolidated[key] = result[key]
                        break
            
            print(f"\n   🎯 RESULTADO CONSOLIDADO:")
            for key, value in consolidated.items():
                status = "✅" if value != "Não encontrado" else "❌"
                print(f"      {status} {key}: {value}")
            
            return consolidated
        else:
            print(f"\n   ⚠️ Nenhum resultado consolidado obtido")
            return self._get_empty_result()
    
    def _get_empty_result(self) -> Dict:
        return {
            'Valor_Global': "Não extraído",
            'Valor_Maximo_Por_Projeto': "Não extraído",
            'Data_Limite_Submissao': "Não extraído",
            'Percentual_Contrapartida': "Não extraído",
            'Nivel_TRL_Exigido': "Não extraído"
        }
    
    def _get_error_result(self, error_msg: str) -> Dict:
        return {
            'Valor_Global': f"Erro: {error_msg[:50]}",
            'Valor_Maximo_Por_Projeto': "Erro na análise",
            'Data_Limite_Submissao': "Erro na análise",
            'Percentual_Contrapartida': "Erro na análise",
            'Nivel_TRL_Exigido': "Erro na análise"
        }


import unittest
from unittest.mock import Mock, patch, MagicMock
import io
from datetime import datetime

from ai.pdf_analyzer import PDFAnalyzer


class TestPDFAnalyzer(unittest.TestCase):
    
    def setUp(self):
        """Configuração inicial dos testes"""
        with patch('ai.pdf_analyzer.OpenAIClient'):
            self.analyzer = PDFAnalyzer()
    
    # ===== TESTES DE INICIALIZAÇÃO =====
    
    def test_initialization_success(self):
        """Testa inicialização bem-sucedida"""
        with patch('ai.pdf_analyzer.OpenAIClient') as mock_client:
            analyzer = PDFAnalyzer()
            self.assertIsNotNone(analyzer.openai_client)
            mock_client.assert_called_once()
    
    # ===== TESTES DE HELPER METHODS =====
    
    def test_guess_ext_pdf(self):
        """Testa extração de extensão .pdf"""
        url = "https://exemplo.com/edital_2025.pdf"
        self.assertEqual(self.analyzer._guess_ext(url), ".pdf")
    
    def test_guess_ext_docx(self):
        """Testa extração de extensão .docx"""
        url = "https://exemplo.com/documento.docx"
        self.assertEqual(self.analyzer._guess_ext(url), ".docx")
    
    def test_guess_ext_uppercase(self):
        """Testa que extensões em maiúsculas são convertidas"""
        url = "https://exemplo.com/arquivo.PDF"
        self.assertEqual(self.analyzer._guess_ext(url), ".pdf")
    
    def test_guess_ext_no_extension(self):
        """Testa URL sem extensão"""
        url = "https://exemplo.com/pagina"
        self.assertEqual(self.analyzer._guess_ext(url), "")
    
    # ===== TESTES DE EXTRAÇÃO DE PDF =====
    
    @patch('ai.pdf_analyzer.fitz')
    def test_extract_text_from_pdf_success(self, mock_fitz):
        """Testa extração bem-sucedida de texto de PDF"""
        mock_page1 = MagicMock()
        mock_page1.get_text.return_value = "Página 1: Edital de Chamada Pública"
        mock_page2 = MagicMock()
        mock_page2.get_text.return_value = "Página 2: Valor Global R$ 1.000.000"
        
        mock_doc = MagicMock()
        mock_doc.__enter__.return_value = [mock_page1, mock_page2]
        mock_fitz.open.return_value = mock_doc
        
        result = self.analyzer._extract_text_from_pdf(b"fake pdf bytes")
        
        self.assertIn("Página 1", result)
        self.assertIn("Página 2", result)
        self.assertIn("Edital", result)
    
    @patch('ai.pdf_analyzer.fitz', None)
    def test_extract_text_from_pdf_no_fitz(self):
        """Testa extração quando PyMuPDF não está disponível"""
        result = self.analyzer._extract_text_from_pdf(b"fake pdf")
        self.assertEqual(result, "")
    
    @patch('ai.pdf_analyzer.fitz')
    def test_extract_text_from_pdf_error(self, mock_fitz):
        """Testa tratamento de erro na extração de PDF"""
        mock_fitz.open.side_effect = Exception("Erro de leitura")
        result = self.analyzer._extract_text_from_pdf(b"corrupted pdf")
        self.assertEqual(result, "")
    
    # ===== TESTES DE EXTRAÇÃO DE DOCX =====
    
    @patch('ai.pdf_analyzer.docx')
    def test_extract_text_from_docx_success(self, mock_docx):
        """Testa extração bem-sucedida de DOCX"""
        mock_para1 = MagicMock()
        mock_para1.text = "Edital Nº 001/2025"
        mock_para2 = MagicMock()
        mock_para2.text = "Valor: R$ 2.000.000"
        
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2]
        mock_docx.Document.return_value = mock_doc
        
        result = self.analyzer._extract_text_from_docx(b"fake docx")
        
        self.assertIn("Edital Nº 001/2025", result)
        self.assertIn("Valor: R$ 2.000.000", result)
    
    @patch('ai.pdf_analyzer.docx', None)
    def test_extract_text_from_docx_no_docx(self):
        """Testa extração quando python-docx não está disponível"""
        result = self.analyzer._extract_text_from_docx(b"fake docx")
        self.assertEqual(result, "")
    
    # ===== TESTES DE EXTRAÇÃO DE HTML =====
    
    @patch('ai.pdf_analyzer.trafilatura')
    def test_extract_text_from_html_with_trafilatura(self, mock_trafilatura):
        """Testa extração com trafilatura"""
        mock_trafilatura.extract.return_value = "Conteúdo limpo extraído"
        
        html = "<html><body><p>Teste</p></body></html>"
        result = self.analyzer._extract_text_from_html(html, "http://exemplo.com")
        
        self.assertEqual(result, "Conteúdo limpo extraído")
    
    @patch('ai.pdf_analyzer.trafilatura', None)
    @patch('ai.pdf_analyzer.BeautifulSoup')
    def test_extract_text_from_html_with_bs4(self, mock_bs):
        """Testa extração com BeautifulSoup"""
        mock_soup = MagicMock()
        mock_soup.get_text.return_value = "Texto extraído com BS4"
        mock_bs.return_value = mock_soup
        
        html = "<html><body><p>Teste</p></body></html>"
        result = self.analyzer._extract_text_from_html(html, "http://exemplo.com")
        
        self.assertEqual(result, "Texto extraído com BS4")
    
    # ===== TESTES DE ANÁLISE DE PDF COM IA =====
    
    def test_analyze_pdf_ai_unavailable(self):
        """Testa análise quando IA não está disponível"""
        self.analyzer.openai_client.is_available = Mock(return_value=False)
        
        result = self.analyzer.analyze_pdf_with_ai_robust("texto qualquer")
        
        self.assertEqual(result['Valor_Global'], "IA não disponível")
        self.assertEqual(result['Valor_Maximo_Por_Projeto'], "IA não disponível")
        self.assertEqual(result['Data_Limite_Submissao'], "IA não disponível")
    
    @patch('ai.pdf_analyzer.split_text_intelligently')
    def test_analyze_pdf_ai_success(self, mock_split):
        """Testa análise bem-sucedida com IA"""
        mock_split.return_value = ["chunk1"]
        self.analyzer.openai_client.is_available = Mock(return_value=True)
        
        mock_response = Mock()
        mock_choice = Mock()
        mock_choice.message.content = '{"Valor_Global": "R$ 5.000.000", "Valor_Maximo_Por_Projeto": "R$ 500.000", "Data_Limite_Submissao": "31/12/2025", "Percentual_Contrapartida": "15%", "Nivel_TRL_Exigido": "TRL 6"}'
        mock_response.choices = [mock_choice]
        
        self.analyzer.openai_client.client.chat.completions.create = Mock(return_value=mock_response)
        
        result = self.analyzer.analyze_pdf_with_ai_robust("Edital completo...")
        
        self.assertEqual(result['Valor_Global'], "R$ 5.000.000")
        self.assertEqual(result['Valor_Maximo_Por_Projeto'], "R$ 500.000")
        self.assertEqual(result['Data_Limite_Submissao'], "31/12/2025")
    
    @patch('ai.pdf_analyzer.split_text_intelligently')
    def test_analyze_pdf_ai_exception(self, mock_split):
        """Testa tratamento de exceção na análise"""
        mock_split.side_effect = Exception("Erro crítico")
        self.analyzer.openai_client.is_available = Mock(return_value=True)
        
        result = self.analyzer.analyze_pdf_with_ai_robust("texto")
        
        self.assertIn("Erro:", result['Valor_Global'])
        self.assertEqual(result['Valor_Maximo_Por_Projeto'], "Erro na análise")
    
    # ===== TESTES DE IDENTIFICAÇÃO DE EDITAL =====
    
    def test_is_pdf_edital_positive_keywords(self):
        """Testa identificação positiva com palavras-chave"""
        text = """
        EDITAL DE CHAMADA PÚBLICA Nº 001/2025
        Objeto: Seleção de propostas de pesquisa
        Valor Global: R$ 3.000.000
        Data de Submissão: 31/03/2025
        Critérios de elegibilidade aplicáveis
        """
        
        is_edital, confidence = self.analyzer.is_pdf_edital(text)
        
        self.assertTrue(is_edital)
        self.assertGreater(confidence, 0.3)
    
    def test_is_pdf_edital_negative_keywords(self):
        """Testa rejeição com palavras-chave negativas"""
        text = """
        Tutorial: Como acessar o sistema
        Guia passo a passo para login
        Manual de instalação do software
        FAQ de suporte técnico
        """
        
        is_edital, confidence = self.analyzer.is_pdf_edital(text)
        
        self.assertFalse(is_edital)
    
    def test_is_pdf_edital_with_project_name_match(self):
        """Testa boost de confiança com nome do projeto"""
        text = "Edital INOVA 2025 - Chamada para inovação tecnológica"
        url = "http://exemplo.com/edital_inova_2025.pdf"
        
        is_edital, confidence = self.analyzer.is_pdf_edital(text, url, "INOVA")
        
        self.assertTrue(is_edital)
        self.assertGreater(confidence, 0.2)
    
    def test_is_pdf_edital_with_current_year(self):
        """Testa boost com ano atual no texto"""
        current_year = str(datetime.utcnow().year)
        text = f"Chamada Pública {current_year} - Fomento à Pesquisa"
        
        is_edital, confidence = self.analyzer.is_pdf_edital(text)
        
        self.assertTrue(is_edital)
    
    @patch('ai.pdf_analyzer.safe_json_parse')
    def test_is_pdf_edital_with_ai_classification(self, mock_parse):
        """Testa classificação usando IA"""
        self.analyzer.openai_client.is_available = Mock(return_value=True)
        
        mock_response = Mock()
        mock_choice = Mock()
        mock_choice.message.content = '{"is_edital": true, "confidence": 0.92}'
        mock_response.choices = [mock_choice]
        
        self.analyzer.openai_client.client.chat.completions.create = Mock(return_value=mock_response)
        mock_parse.return_value = {"is_edital": True, "confidence": 0.92}
        
        long_text = "Edital completo com regras detalhadas " * 50
        is_edital, confidence = self.analyzer.is_pdf_edital(long_text, "http://exemplo.com/edital.pdf")
        
        self.assertTrue(is_edital)
        self.assertGreater(confidence, 0.5)
    
    # ===== TESTES DE MÉTODOS PRIVADOS =====
    
    def test_create_analysis_prompt(self):
        """Testa criação de prompt de análise"""
        chunk = "Texto exemplo do edital"
        prompt = self.analyzer._create_analysis_prompt(chunk)
        
        self.assertIn("MISSÃO ESPECIALIZADA", prompt)
        self.assertIn("Valor Global", prompt)
        self.assertIn("Valor Máximo do Projeto", prompt)
        self.assertIn(chunk, prompt)
    
    def test_parse_ai_response_complete(self):
        """Testa parse de resposta completa da IA"""
        response = '''
        {
            "Valor_Global": "R$ 10.000.000",
            "Valor_Maximo_Por_Projeto": "R$ 1.000.000",
            "Data_Limite_Submissao": "15/06/2025",
            "Percentual_Contrapartida": "20%",
            "Nivel_TRL_Exigido": "TRL 7-8"
        }
        '''
        
        result = self.analyzer._parse_ai_response(response)
        
        self.assertEqual(result['Valor_Global'], "R$ 10.000.000")
        self.assertEqual(result['Valor_Maximo_Por_Projeto'], "R$ 1.000.000")
        self.assertEqual(result['Data_Limite_Submissao'], "15/06/2025")
        self.assertEqual(result['Percentual_Contrapartida'], "20%")
        self.assertEqual(result['Nivel_TRL_Exigido'], "TRL 7-8")
    
    def test_parse_ai_response_partial(self):
        """Testa parse de resposta parcial"""
        response = '{"Valor_Global": "R$ 5.000.000"}'
        
        result = self.analyzer._parse_ai_response(response)
        
        self.assertEqual(result['Valor_Global'], "R$ 5.000.000")
        self.assertEqual(result['Valor_Maximo_Por_Projeto'], "Não encontrado")
    
    def test_extract_found_info(self):
        """Testa extração de informações encontradas"""
        result = {
            'Valor_Global': "R$ 2.000.000",
            'Valor_Maximo_Por_Projeto': "Não encontrado",
            'Data_Limite_Submissao': "30/04/2025",
            'Percentual_Contrapartida': "10%",
            'Nivel_TRL_Exigido': "Não encontrado"
        }
        
        found = self.analyzer._extract_found_info(result)
        
        self.assertEqual(len(found), 3)
        self.assertIn("Valor_Global: R$ 2.000.000", found)
        self.assertIn("Data_Limite_Submissao: 30/04/2025", found)
        self.assertIn("Percentual_Contrapartida: 10%", found)
    
    def test_consolidate_results_multiple_chunks(self):
        """Testa consolidação de múltiplos chunks"""
        results = [
            {
                'Valor_Global': "R$ 8.000.000",
                'Valor_Maximo_Por_Projeto': "Não encontrado",
                'Data_Limite_Submissao': "Não encontrado",
                'Percentual_Contrapartida': "Não encontrado",
                'Nivel_TRL_Exigido': "Não encontrado"
            },
            {
                'Valor_Global': "Não encontrado",
                'Valor_Maximo_Por_Projeto': "R$ 800.000",
                'Data_Limite_Submissao': "20/05/2025",
                'Percentual_Contrapartida': "Não encontrado",
                'Nivel_TRL_Exigido': "Não encontrado"
            },
            {
                'Valor_Global': "Não encontrado",
                'Valor_Maximo_Por_Projeto': "Não encontrado",
                'Data_Limite_Submissao': "Não encontrado",
                'Percentual_Contrapartida': "12%",
                'Nivel_TRL_Exigido': "TRL 5"
            }
        ]
        
        consolidated = self.analyzer._consolidate_results(results)
        
        self.assertEqual(consolidated['Valor_Global'], "R$ 8.000.000")
        self.assertEqual(consolidated['Valor_Maximo_Por_Projeto'], "R$ 800.000")
        self.assertEqual(consolidated['Data_Limite_Submissao'], "20/05/2025")
        self.assertEqual(consolidated['Percentual_Contrapartida'], "12%")
        self.assertEqual(consolidated['Nivel_TRL_Exigido'], "TRL 5")
    
    def test_consolidate_results_empty(self):
        """Testa consolidação com lista vazia"""
        consolidated = self.analyzer._consolidate_results([])
        
        for value in consolidated.values():
            self.assertEqual(value, "Não extraído")
    
    def test_get_empty_result(self):
        """Testa estrutura de resultado vazio"""
        result = self.analyzer._get_empty_result()
        
        self.assertIn('Valor_Global', result)
        self.assertIn('Valor_Maximo_Por_Projeto', result)
        self.assertEqual(result['Valor_Global'], "Não extraído")
    
    def test_get_error_result(self):
        """Testa estrutura de resultado com erro"""
        error_msg = "Erro ao processar documento: timeout na requisição"
        result = self.analyzer._get_error_result(error_msg)
        
        self.assertIn("Erro:", result['Valor_Global'])
        self.assertEqual(result['Valor_Maximo_Por_Projeto'], "Erro na análise")
        self.assertLessEqual(len(result['Valor_Global']), 56)  # Truncado


if __name__ == '__main__':
    # Executa os testes
    unittest.main(verbosity=2)